import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def create_similar_document():
    doc = Document()

    # Set margins to 0.9 inches on all sides (matching the reference document)
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Color Palette Constants (Matching Reference Document)
    COLOR_PRIMARY = RGBColor(26, 86, 219)     # Royal Blue (#1A56DB)
    COLOR_DARK_TEXT = RGBColor(17, 24, 39)    # Dark Charcoal (#111827)
    COLOR_BODY_TEXT = RGBColor(55, 65, 81)    # Dark Slate Gray (#374151)
    COLOR_MUTED = RGBColor(107, 114, 128)     # Muted Gray (#6B7280)
    
    # Table styling hex values
    HEX_PRIMARY = "1A56DB"
    HEX_LIGHT_BG = "F9FAFB"                   # Light Gray 50 (#F9FAFB)
    HEX_BORDER = "E5E7EB"                      # Light Border (#E5E7EB)

    # ---------------------------------------------------------
    # Helper Functions for Formatting
    # ---------------------------------------------------------
    def set_font_run(run, name="Calibri", size_pt=9.5, color_rgb=COLOR_BODY_TEXT, bold=False, italic=False):
        run.font.name = name
        run.font.size = Pt(size_pt)
        if color_rgb:
            run.font.color.rgb = color_rgb
        run.bold = bold
        run.italic = italic

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        
        borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
        for name, b_style in borders.items():
            if b_style is not None:
                val = b_style.get('val', 'single')
                sz = b_style.get('sz', '4')
                space = b_style.get('space', '0')
                color = b_style.get('color', 'auto')
                el = parse_xml(f'<w:{name} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="{space}" w:color="{color}"/>')
                tcBorders.append(el)
            else:
                el = parse_xml(f'<w:{name} {nsdecls("w")} w:val="none"/>')
                tcBorders.append(el)
        tcPr.append(tcBorders)

    def add_custom_heading(number_text, title_text, space_before=16, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        
        # Heading numbers are in COLOR_PRIMARY
        r_num = p.add_run(number_text)
        set_font_run(r_num, size_pt=13, color_rgb=COLOR_PRIMARY, bold=True)
        
        # Heading titles are in COLOR_DARK_TEXT
        r_title = p.add_run(title_text)
        set_font_run(r_title, size_pt=13, color_rgb=COLOR_DARK_TEXT, bold=True)
        return p

    def add_custom_subheading(text, space_before=12, space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text)
        set_font_run(run, size_pt=11, color_rgb=COLOR_DARK_TEXT, bold=True)
        return p

    def add_usecase_title(text, space_before=10, space_after=2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text)
        set_font_run(run, size_pt=10, color_rgb=COLOR_BODY_TEXT, bold=True, italic=True)
        return p

    def add_custom_paragraph(text="", space_after=6, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if text:
            run = p.add_run(text)
            set_font_run(run, italic=italic)
        return p

    def add_custom_list_item(text, space_after=3):
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        set_font_run(run)
        return p

    # ---------------------------------------------------------
    # COVER BLOCK (Matching Reference Page 1 Title block)
    # ---------------------------------------------------------
    # P0: Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("AMC Research Portal")
    set_font_run(run_title, size_pt=24, color_rgb=COLOR_PRIMARY, bold=True)
    
    # P1: Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(2)
    run_sub = p_sub.add_run("AI Initiative & Use Case Document")
    set_font_run(run_sub, size_pt=12, color_rgb=COLOR_BODY_TEXT)

    # P2: Subsubtitle
    p_subsub = doc.add_paragraph()
    p_subsub.paragraph_format.space_after = Pt(16)
    run_subsub = p_subsub.add_run("Institutional Event-Intelligence & Monitoring System — Nifty 500")
    set_font_run(run_subsub, size_pt=9.5, color_rgb=COLOR_MUTED, italic=True)

    # ---------------------------------------------------------
    # SECTION 1: AI Initiative & Use Case Identification
    # ---------------------------------------------------------
    add_custom_heading("1  ", "AI Initiative & Use Case Identification")
    
    add_custom_paragraph(
        "AMC Research Portal is an institutional-grade event-intelligence and monitoring system for Indian equity "
        "markets (Nifty 500 constituents) that helps fund managers and research analysts track high-impact corporate "
        "developments. The AI initiative embeds multi-agent intelligence at every layer of the product — from real-time "
        "corporate filing ingestion and verification to sentiment scoring and historical analog analysis — to accelerate "
        "decision-making, reduce risk exposure, and drive measurable performance alpha."
    )
    
    add_custom_subheading("Core AI Use Cases")
    
    # UC-01
    add_usecase_title("UC-01 · Ingestion & Constituent Resolution")
    add_custom_list_item("A background daemon dynamically parses the semi-annual Nifty 500 constituent listings from the NSE portal, resolving ticker mappings and establishing the active monitoring universe.")
    add_custom_list_item("Automated scraping of exchange announcements (BSE & NSE feeds) is anchored to the actual calendar year/date to bypass model temporal drift and ensure data freshness.")
    add_custom_list_item("Output: A structured watch list mapping events to specific tickers, avoiding penny stocks and non-constituent entities.")

    # UC-02
    add_usecase_title("UC-02 · Automated Event Verification & Deduplication")
    add_custom_list_item("Ensures high-severity developments (e.g., senior executive resignations, forensic audits, debt defaults) are cross-referenced across official Tier 1 filings and Tier 2 trusted financial press (Reuters, Bloomberg, Moneycontrol).")
    add_custom_list_item("Deduplicates announcements sharing the same company, event type, and date, merging disparate media articles into a single coherent event card.")
    add_custom_list_item("System path: If an event cannot be verified against official exchange filings, the agent abstains from scoring and logs it under 'Insufficient Verification'.")

    # UC-03
    add_usecase_title("UC-03 · Actionability Scoring & Impact Briefs")
    add_custom_list_item("A multi-agent consensus pipeline assigns a composite Actionability Score (0–100) based on intrinsic severity (nature of the event), expected fundamental magnitude (balance sheet impact), and source confidence.")
    add_custom_list_item("Generates structured Executive Briefs with actionable summaries, directional leans, and clear inline links to primary sources, bypassing long PDF filings.")

    # UC-04
    add_usecase_title("UC-04 · Contextual Historical Analog Engine")
    add_custom_list_item("Searches the database for comparable historical events (e.g., historical auditor resignations in the same sector) and calculates similarity scores.")
    add_custom_list_item("Retrieves realized return profiles (1-day, 5-day, 30-day post-event) of those analogs to establish a statistical baseline for market reaction.")
    add_custom_list_item("Enables fund managers to base decisions on historical empirical evidence rather than simple subjective assessment.")

    # ---------------------------------------------------------
    # TABLE 0: Target Segment and Platform Edge
    # ---------------------------------------------------------
    table0 = doc.add_table(rows=5, cols=2)
    table0.alignment = WD_TABLE_ALIGNMENT.CENTER
    table0.autofit = False
    
    t0_widths = [Inches(1.8), Inches(4.7)]
    t0_data = [
        ("Target Segment", "Institutional fund managers, equity research analysts, and investment committee members managing Indian equities."),
        ("Market Opportunity", "Nifty 500 represents 90%+ of free-float cap; manual processing of BSE/NSE corporate disclosures results in delayed trading responses."),
        ("Platform Edge", "Only portal integrating real-time corporate filing scrapers, dual-channel verification check, and automated historical analog back-tests."),
        ("AI Innovation Goal", "Reduce time-to-first-risk-response on negative announcements from hours/days to ≤ 15 minutes of exchange publication."),
        ("Efficiency Signal", "Reduction in analyst filtration time; higher accuracy in flagging severe corporate governance events before full market absorption.")
    ]
    
    border_style = {'val': 'single', 'sz': '4', 'color': HEX_BORDER}
    
    for row_idx, (cat, desc) in enumerate(t0_data):
        row = table0.rows[row_idx]
        
        # Col 0 (Category)
        cell0 = row.cells[0]
        cell0.width = t0_widths[0]
        cell0.text = cat
        set_cell_background(cell0, HEX_LIGHT_BG)
        p0 = cell0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after = Pt(4)
        set_font_run(p0.runs[0], size_pt=9.5, color_rgb=COLOR_DARK_TEXT, bold=True)
        set_cell_borders(cell0, top=border_style, bottom=border_style, left=border_style, right=border_style)
        
        # Col 1 (Description)
        cell1 = row.cells[1]
        cell1.width = t0_widths[1]
        cell1.text = desc
        p1 = cell1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(4)
        p1.paragraph_format.line_spacing = 1.15
        set_font_run(p1.runs[0], size_pt=9.5, color_rgb=COLOR_BODY_TEXT)
        set_cell_borders(cell1, top=border_style, bottom=border_style, left=border_style, right=border_style)

    add_custom_paragraph() # Spacer

    # ---------------------------------------------------------
    # SECTION 2: Strategic Opportunity & AI Innovation Targets
    # ---------------------------------------------------------
    add_custom_heading("2  ", "Strategic Opportunity & AI Innovation Targets")
    
    add_custom_paragraph(
        "Indian equity markets are experiencing structural formalisation, with domestic institutional inflows hitting "
        "record highs. For fund managers, capturing alpha is increasingly dependent on speed and precision. Information "
        "overload from corporate actions and regulatory changes acts as a bottleneck. AMC Research Portal addresses this "
        "gap by automating the heavy lifting of document search, multi-source verification, and historical grounding — "
        "letting analysts focus on valuation and strategy."
    )
    
    add_custom_subheading("Innovation Targets (6-Month Horizon)")
    add_custom_list_item("Automated drafting of investment committee compliance and research notes from verified event briefs.")
    add_custom_list_item("Integration of natural language query capabilities allowing fund managers to ask questions about current portfolio risk exposure to specific events.")
    add_custom_list_item("Automatic text and sentiment analysis of earnings call transcripts, mapping executive vocal distress or tone anomalies.")
    add_custom_list_item("Cross-asset correlation mapping: connecting equity event alerts with corresponding debt/credit rating changes and currency fluctuations.")

    # ---------------------------------------------------------
    # SECTION 3: Value Realization Framework: Strategy & AI Feasibility Metrics
    # ---------------------------------------------------------
    add_custom_heading("3  ", "Value Realization Framework: Strategy & AI Feasibility Metrics")
    
    add_custom_subheading("Feasibility Assessment")
    add_custom_list_item("\u2705  SQLite database (dev.db) schema supporting User, Profile, AgentRun, and EventState tables is live.")
    add_custom_list_item("\u2705  Multi-LLM provider integration (Gemini, OpenAI, Claude) successfully implemented in the backend portal.")
    add_custom_list_item("\u2705  NextJS-based UI layout, profile filtering settings, and user authentication are active.")
    add_custom_list_item("\u2705  Scraper architecture with temporal anchoring and date-year constraints validated.")
    add_custom_list_item("\u26a0  Live WebSocket filings feed and advanced news API integration are planned additions (Q3).")

    # Table 1: Value Realization Metrics
    table1 = doc.add_table(rows=8, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = False
    
    t1_widths = [Inches(1.8), Inches(1.3), Inches(1.3), Inches(2.1)]
    t1_headers = ["Metric", "Baseline (Pre-AI)", "AI Target", "Measurement Method"]
    t1_data = [
        ("Daily Active Users", "—", "+40% MoM", "User session logs in app.User; profile alert clicks"),
        ("Ingestion-to-Alert Latency", "> 45 min (manual)", "\u2264 15 minutes", "AgentRun.startedAt vs publication timestamp"),
        ("Event Ingestion Recall", "88% (estimated)", "\u2265 99% of Nifty 500", "Comparison of AgentRun outputs with daily official exchange bulletins"),
        ("Extraction Precision", "—", "\u2265 97% accuracy", "Programmatic validation of JSON variables against manual analyst audits"),
        ("Alert Actionability CTR", "—", "\u2265 45% click-through", "FundManagerProfile.alertScope match count vs review actions taken"),
        ("Downside Alpha Protection", "—", "Exit/trim before >2% drop", "Realized return tracking on event-based portfolio actions"),
        ("Abstention Rate Alignment", "0% (hallucinated)", "100% correct log", "Audit of 'Insufficient Verification' logs vs false-positive alerts")
    ]
    
    # Header Row Styling
    hdr_cells = table1.rows[0].cells
    for col_idx, header_text in enumerate(t1_headers):
        cell = hdr_cells[col_idx]
        cell.width = t1_widths[col_idx]
        cell.text = header_text
        set_cell_background(cell, HEX_PRIMARY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_font_run(p.runs[0], size_pt=9.5, color_rgb=RGBColor(255, 255, 255), bold=True)
        set_cell_borders(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
        
    for row_idx, row_data in enumerate(t1_data, 1):
        row = table1.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = t1_widths[col_idx]
            cell.text = text
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            # Highlight first cell in row
            if col_idx == 0:
                set_font_run(p.runs[0], size_pt=9.5, color_rgb=COLOR_DARK_TEXT, bold=True)
            else:
                set_font_run(p.runs[0], size_pt=9.5, color_rgb=COLOR_BODY_TEXT)
                
            if row_idx % 2 == 1:
                set_cell_background(cell, HEX_LIGHT_BG)
                
            set_cell_borders(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)

    add_custom_paragraph() # Spacer

    # ---------------------------------------------------------
    # SECTION 4: Target Operational Workflow: End-to-End Happy Path
    # ---------------------------------------------------------
    add_custom_heading("4  ", "Target Operational Workflow: End-to-End Happy Path")
    
    add_custom_subheading("System/Admin: Automated Event Scan & Alert Broadcast")
    add_custom_list_item("Step 1 — Background cron task triggers every 15 minutes during market hours, downloading the Nifty 500 index constituent list.")
    add_custom_list_item("Step 2 — Ingestion scraper pulls announcements from BSE/NSE corporate disclosures and searches google news using date-anchored queries.")
    add_custom_list_item("Step 3 — System runs deduplication and verification check. If verified, triggers the multi-agent consensus pipeline for severity and magnitude scoring.")
    add_custom_list_item("Step 4 — Agent runs search queries against the database to map historical analogs and calculate return distributions.")
    add_custom_list_item("Step 5 — Compiled event brief is written to the database (Prisma client) and pushed as high-priority alert to fund managers with matching profile watchlists.")

    add_custom_subheading("Fund Manager: Reviewing Brief & Executing Portfolio Decision")
    add_custom_list_item("Step 1 — Fund manager receives secure email/Slack alert: \"CRITICAL EVENT: auditor resignation at [Company] — Severity 5\".")
    add_custom_list_item("Step 2 — Fund manager opens dashboard; UI loads the interactive 'Executive Brief' detailing the core event and expected margin impact.")
    add_custom_list_item("Step 3 — Manager reviews the Historical Analog section, noting that 4 of 5 past analogs experienced a >5% drop within the first 5 sessions.")
    add_custom_list_item("Step 4 — Manager clicks the primary exchange document link to confirm, then opens trade execution portal to hedge or trim position.")
    add_custom_list_item("Step 5 — Position updated; manager updates EventState status to 'Completed' and submits feedback on scoring accuracy.")

    # ---------------------------------------------------------
    # SECTION 5: AI Operational Scope & Innovation Lifecycle
    # ---------------------------------------------------------
    add_custom_heading("5  ", "AI Operational Scope & Innovation Lifecycle")
    
    add_custom_subheading("5.1  In-Product AI Intervention (Runtime UX & Feature Autonomy)")
    add_custom_paragraph("These AI capabilities run live in production and interact directly with end-users in real time.")
    
    add_custom_paragraph(
        "Model selection rationale: Gemini 3.5 Flash is used for high-frequency ingestion, entity extraction, "
        "and initial classification to keep API latency and token costs low. Gemini 1.5 Pro (or Claude Opus) "
        "is reserved for the Multi-Agent Consensus Scoring and Historical Analog mapping — where reasoning quality, "
        "data integration, and detail are paramount."
    )

    # Table 2: Feature Model & Trigger Mapping
    table2 = doc.add_table(rows=7, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.autofit = False
    
    t2_widths = [Inches(1.8), Inches(1.2), Inches(1.5), Inches(2.0)]
    t2_headers = ["Feature", "Model", "Trigger", "Output"]
    t2_data = [
        ("Constituent Resolution", "JSON parser (rules)", "Scheduled 08:30 IST cron", "Active Nifty 500 constituent list with mapped sector tickers"),
        ("Event Ingestion Scrape", "Scraping modules", "15-minute cron interval", "Raw corporate announcement files and financial news articles"),
        ("Multi-Source Verification", "Gemini 3.5 Flash", "Ingestion file discovery", "Deduplicated, validated event records or 'Insufficient Verification' logs"),
        ("Actionability Consensus", "Gemini 1.5 Pro", "Post-verification trigger", "Structured JSON with Severity (1-5), Magnitude (1-5), and expected timeline"),
        ("Historical Analog Mapping", "Database + LLM", "Post-scoring trigger", "Top 3 verifiable past analogs with return curves and similarity scores"),
        ("Alert Push Dispatch", "Prisma & Server push", "High Actionability Event", "Real-time Slack, email, or in-app notification to matched managers")
    ]
    
    # Header Row Styling
    hdr_cells2 = table2.rows[0].cells
    for col_idx, header_text in enumerate(t2_headers):
        cell = hdr_cells2[col_idx]
        cell.width = t2_widths[col_idx]
        cell.text = header_text
        set_cell_background(cell, HEX_PRIMARY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_font_run(p.runs[0], size_pt=9.5, color_rgb=RGBColor(255, 255, 255), bold=True)
        set_cell_borders(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
        
    for row_idx, row_data in enumerate(t2_data, 1):
        row = table2.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = t2_widths[col_idx]
            cell.text = text
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            if col_idx == 0:
                set_font_run(p.runs[0], size_pt=9.5, color_rgb=COLOR_DARK_TEXT, bold=True)
            else:
                set_font_run(p.runs[0], size_pt=9.5, color_rgb=COLOR_BODY_TEXT)
                
            if row_idx % 2 == 1:
                set_cell_background(cell, HEX_LIGHT_BG)
                
            set_cell_borders(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)

    add_custom_paragraph() # Spacer

    add_custom_subheading("5.2  In-Development AI Assistance (Engineering, Prototyping & Training Phase)")
    add_custom_paragraph("AI is also embedded in the development lifecycle to accelerate feature delivery and improve model output quality.")
    
    add_custom_paragraph(
        "All AI model calls are instrumented with token counts, latency, and error rates. LLM responses are validated "
        "against a strict JSON schema before persisting — invalid outputs fall back to default rules-based logs so that "
        "the user experience is never blocked by model failure."
    )

    # Table 3: In-Development Phase Mapping
    table3 = doc.add_table(rows=7, cols=4)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.autofit = False
    
    t3_widths = [Inches(1.5), Inches(1.5), Inches(1.5), Inches(2.0)]
    t3_headers = ["Phase", "AI Role", "Tool / Method", "Outcome"]
    t3_data = [
        ("Prompt Engineering", "Iterate on system prompts for extraction & analog matching", "Prompt versioning in Git", "Higher extraction precision; fewer hallucinated dates/numbers"),
        ("Regression Testing", "Evaluate prompt changes against historical gold-standard set", "Automated prompt runner script", "Prevent degradation of event categorization accuracy before deployment"),
        ("Sandbox Seeding", "Generate realistic event arrays for local testing", "seedEvents.js script", "Consistent mock data in local dev.db without calling live APIs"),
        ("Scraper Simulation", "Simulate network errors, rate limits, and PDF formatting changes", "Mock scraping harness", "Robust error-handling and fallback logic in crawler modules"),
        ("Feedback Loop Export", "Extract user-corrected ratings and scores to dataset", "Database export to JSON", "Dataset ready for prompt tuning and potential model fine-tuning"),
        ("Documentation", "Auto-generate API documentation and database schema diagrams", "AI-assisted doc generator", "Up-to-date developer docs matching the schema.prisma changes")
    ]
    
    # Header Row Styling
    hdr_cells3 = table3.rows[0].cells
    for col_idx, header_text in enumerate(t3_headers):
        cell = hdr_cells3[col_idx]
        cell.width = t3_widths[col_idx]
        cell.text = header_text
        set_cell_background(cell, HEX_PRIMARY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_font_run(p.runs[0], size_pt=9.5, color_rgb=RGBColor(255, 255, 255), bold=True)
        set_cell_borders(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
        
    for row_idx, row_data in enumerate(t3_data, 1):
        row = table3.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = t3_widths[col_idx]
            cell.text = text
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            if col_idx == 0:
                set_font_run(p.runs[0], size_pt=9.5, color_rgb=COLOR_DARK_TEXT, bold=True)
            else:
                set_font_run(p.runs[0], size_pt=9.5, color_rgb=COLOR_BODY_TEXT)
                
            if row_idx % 2 == 1:
                set_cell_background(cell, HEX_LIGHT_BG)
                
            set_cell_borders(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)

    # ---------------------------------------------------------
    # FOOTER SETUP
    # ---------------------------------------------------------
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("AMC Research Portal  \u00b7  Confidential  \u00b7  15 July 2026")
    set_font_run(footer_run, name="Calibri", size_pt=9.5, color_rgb=COLOR_MUTED)

    # Save the document
    filename = "AMC_Research_Portal_AI_UseCase.docx"
    doc.save(filename)
    print(f"Document successfully created and saved as {filename}")

if __name__ == "__main__":
    create_similar_document()
